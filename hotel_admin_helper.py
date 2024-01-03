import datetime
import pickle
import wx
import wx.adv
import os
import docx
from num2words import num2words
from openpyxl import load_workbook
from decimal import Decimal
from enum import Enum
from dataclasses import dataclass, asdict
from docx2pdf import convert


# constants
TODAY = datetime.date.today()
APP_EXIT = wx.NewIdRef()
APP_SAVE = wx.NewIdRef()
CHANGE_PRICES_FRAME = wx.NewIdRef()
TWOPLACES = Decimal(10) ** -2
app = wx.App()
curr_path = os.getcwd()


# main frame
class MyFrame(wx.Frame):
    def __init__(self, parent, title):
        super().__init__(parent, title=title)

        # Menu_bar
        self.content = None
        menubar = wx.MenuBar()
        # File Menu
        fileMenu = wx.Menu()

        fileMenu.AppendSeparator()  # Separator

        exit_menu = wx.MenuItem(fileMenu, APP_EXIT, "Вихід\tCtrl+Q", "Вийти з додатку")
        fileMenu.Append(exit_menu)

        menubar.Append(fileMenu, "&Файл")

        # Settings menu
        settings_menu = wx.Menu()
        settings_price = wx.MenuItem(settings_menu, CHANGE_PRICES_FRAME, "&Ціни\tCtrl+P", "Встанови ціни "
                                                                                          "номери/тур.збір/Сніданок")
        settings_menu.Append(settings_price)

        menubar.Append(settings_menu, "&Налаштування")

        # menu binds
        self.SetMenuBar(menubar)
        self.Bind(wx.EVT_MENU, self.onquit, id=APP_EXIT)
        self.Bind(wx.EVT_MENU, self.show_settings_price_frame, id=CHANGE_PRICES_FRAME)

        # main elements
        panel = wx.Panel(self)

        main_sizer = wx.GridBagSizer(10, 10)

        button_confirm = wx.Button(panel, label="Створити підтвердження")
        button_bill = wx.Button(panel, label="Створити рахунок")
        button_bill_wc = wx.Button(panel, label="Створити рахунок безготівковий")
        button_act = wx.Button(panel, label="Створити акт")

        main_sizer.Add(button_confirm, pos=(0, 0), flag=wx.EXPAND | wx.LEFT | wx.TOP, border=2)
        main_sizer.Add(button_bill, pos=(0, 1), flag=wx.EXPAND | wx.LEFT | wx.TOP, border=2)
        main_sizer.Add(button_bill_wc, pos=(0, 2), flag=wx.EXPAND | wx.LEFT | wx.TOP, border=2)
        main_sizer.Add(button_act, pos=(0, 3), flag=wx.EXPAND | wx.LEFT | wx.TOP | wx.RIGHT, border=2)
        main_sizer.AddGrowableCol(0)
        main_sizer.AddGrowableCol(1)
        main_sizer.AddGrowableCol(2)
        main_sizer.AddGrowableCol(3)

        # binds
        button_confirm.Bind(wx.EVT_BUTTON, self.make_confirm, button_confirm)
        button_bill.Bind(wx.EVT_BUTTON, self.make_bill, button_bill)
        button_bill_wc.Bind(wx.EVT_BUTTON, self.make_bill_wc, button_bill_wc)
        button_act.Bind(wx.EVT_BUTTON, self.make_act, button_act)

        # mandatory

        prices_default.load_from_file("prices_default.pkl")

        guest_name_stat_txt = wx.StaticText(panel, label="Ім'я гостя:")
        main_sizer.Add(guest_name_stat_txt, pos=(1, 0), flag=wx.LEFT, border=10)
        self.guest_name_text_ctrl = wx.TextCtrl(panel)
        main_sizer.Add(self.guest_name_text_ctrl, pos=(1, 1), flag=wx.EXPAND | wx.LEFT, border=10)

        date_make_stat_txt = wx.StaticText(panel, label="Дата формування:")
        main_sizer.Add(date_make_stat_txt, pos=(2, 0), flag=wx.LEFT, border=10)
        self.date_make = wx.adv.DatePickerCtrl(panel, wx.ID_ANY, wx.DefaultDateTime,
                                               style=wx.adv.DP_DROPDOWN | wx.adv.DP_SHOWCENTURY)
        self.date_make.Bind(wx.adv.EVT_DATE_CHANGED, self.make_date_changed)
        main_sizer.Add(self.date_make, pos=(2, 1), flag=wx.EXPAND | wx.LEFT, border=10)

        checkin_date_stat_txt = wx.StaticText(panel, label="CheckIn дата:")
        main_sizer.Add(checkin_date_stat_txt, pos=(3, 0), flag=wx.LEFT, border=10)
        self.checkin_date = wx.adv.DatePickerCtrl(panel, wx.ID_ANY, wx.DefaultDateTime,
                                                  style=wx.adv.DP_DROPDOWN | wx.adv.DP_SHOWCENTURY)
        self.checkin_date.Bind(wx.adv.EVT_DATE_CHANGED, self.checkin_date_changed)
        main_sizer.Add(self.checkin_date, pos=(3, 1), flag=wx.EXPAND | wx.LEFT, border=10)

        checkout_date_stat_txt = wx.StaticText(panel, label="CheckOut дата:")
        main_sizer.Add(checkout_date_stat_txt, pos=(4, 0), flag=wx.LEFT, border=10)
        self.checkout_date = wx.adv.DatePickerCtrl(panel, wx.ID_ANY, wx.DefaultDateTime,
                                                   style=wx.adv.DP_DROPDOWN | wx.adv.DP_SHOWCENTURY)
        self.checkout_date.Bind(wx.adv.EVT_DATE_CHANGED, self.checkout_date_changed)
        main_sizer.Add(self.checkout_date, pos=(4, 1), flag=wx.EXPAND | wx.LEFT, border=10)
        self.tomorrow = TODAY + datetime.timedelta(days=1)
        self.checkout_date.SetValue(self.tomorrow)

        category_stat_txt = wx.StaticText(panel, label="Категорія номеру:")
        main_sizer.Add(category_stat_txt, pos=(5, 0), flag=wx.LEFT, border=10)
        categories = [room_category for room_category in RoomType]
        self.category = wx.ComboBox(panel, choices=categories, style=wx.CB_READONLY)
        main_sizer.Add(self.category, pos=(5, 1), flag=wx.EXPAND | wx.LEFT, border=10)
        self.category.Bind(wx.EVT_COMBOBOX, self.category_combobox)
        default_category = RoomType.Standart
        self.category.SetValue(default_category)

        price_accomodation_PN_stat_txt = wx.StaticText(panel, label="Ціна за добу:")
        main_sizer.Add(price_accomodation_PN_stat_txt, pos=(6, 0), flag=wx.LEFT, border=10)
        self.price_accomodation_PN_text_ctrl = wx.TextCtrl(panel)
        main_sizer.Add(self.price_accomodation_PN_text_ctrl, pos=(6, 1), flag=wx.EXPAND | wx.LEFT, border=10)
        self.price_accomodation_PN_text_ctrl.Bind(wx.EVT_TEXT, self.price_pernight_changed,
                                                  self.price_accomodation_PN_text_ctrl)
        self.price_accomodation_PN_text_ctrl.SetValue(prices_default.prices[self.category.GetValue()])

        total_price_accomodation_stat_txt = wx.StaticText(panel, label="Проживання загальна ціна:")  # float auto-score
        main_sizer.Add(total_price_accomodation_stat_txt, pos=(7, 0), flag=wx.LEFT, border=10)
        self.total_price_accomodation_text_ctrl = wx.TextCtrl(panel)
        main_sizer.Add(self.total_price_accomodation_text_ctrl, pos=(7, 1), flag=wx.EXPAND | wx.LEFT, border=10)

        count_of_guest_stat_txt = wx.StaticText(panel, label="Кількість гостей:")
        main_sizer.Add(count_of_guest_stat_txt, pos=(8, 0), flag=wx.LEFT, border=10)
        count_of_guests = ["1", "2", "3", "4", "5", "6", "7", "8", "9", "10"]
        self.count_of_guest = wx.ComboBox(panel, choices=count_of_guests, style=wx.CB_READONLY)
        main_sizer.Add(self.count_of_guest, pos=(8, 1), flag=wx.EXPAND | wx.LEFT, border=10)
        self.count_of_guest.SetValue(count_of_guests[0])
        self.count_of_guest.Bind(wx.EVT_COMBOBOX, self.count_of_guest_combobox, self.count_of_guest)
        self.count_of_guest.Bind(wx.EVT_COMBOBOX, self.tour_tax_total, self.count_of_guest)

        count_of_rooms_stat_txt = wx.StaticText(panel, label="Кількість номерів:")  # combobox
        main_sizer.Add(count_of_rooms_stat_txt, pos=(9, 0), flag=wx.LEFT, border=10)
        self.count_of_rooms_list = ["1", "2", "3", "4", "5", "6", "7", "8", "9", "10"]
        self.count_of_rooms = wx.ComboBox(panel, choices=self.count_of_rooms_list, style=wx.CB_READONLY)
        main_sizer.Add(self.count_of_rooms, pos=(9, 1), flag=wx.EXPAND | wx.LEFT, border=10)
        self.count_of_rooms.Bind(wx.EVT_COMBOBOX, self.count_of_rooms_combobox, self.count_of_rooms)
        self.count_of_rooms.Bind(wx.EVT_COMBOBOX, self.total_price_accomodation, self.count_of_rooms)
        self.count_of_rooms.SetValue("1")
        self.count_of_rooms.Disable()
        self.count_of_rooms_checkbox = wx.CheckBox(panel)
        main_sizer.Add(self.count_of_rooms_checkbox, pos=(9, 2), flag=wx.ALL, border=5)
        self.count_of_rooms_checkbox.Bind(wx.EVT_CHECKBOX, self.checkbox_count_of_rooms, self.count_of_rooms_checkbox)
        self.count_of_rooms_checkbox.SetValue(False)
        self.total_price_accomodation_text_ctrl.SetValue(str(self.total_price_accomodation(wx.EVT_COMBOBOX)))

        tour_tax_stat_txt = wx.StaticText(panel, label="Туристичний збір:")
        main_sizer.Add(tour_tax_stat_txt, pos=(10, 0), flag=wx.LEFT, border=10)
        self.tour_tax_text_ctrl = wx.TextCtrl(panel)
        main_sizer.Add(self.tour_tax_text_ctrl, pos=(10, 1), flag=wx.EXPAND | wx.LEFT, border=10)
        self.tour_tax_checkbox = wx.CheckBox(panel)
        main_sizer.Add(self.tour_tax_checkbox, pos=(10, 2), flag=wx.ALL, border=5)
        self.tour_tax_checkbox.Bind(wx.EVT_CHECKBOX, self.checkbox_tour_tax, self.tour_tax_checkbox)
        self.tour_tax_checkbox.SetValue(True)

        breakfest_count_stat_txt = wx.StaticText(panel, label="Кількість сніданків:")
        main_sizer.Add(breakfest_count_stat_txt, pos=(11, 0), flag=wx.LEFT, border=10)
        self.breakfest_count_list = ["1", "2", "3", "4", "5", "6", "7", "8", "9", "10", ""]
        self.breakfest_count_combobox = wx.ComboBox(panel, choices=self.breakfest_count_list, style=wx.CB_READONLY)
        main_sizer.Add(self.breakfest_count_combobox, pos=(11, 1), flag=wx.EXPAND | wx.LEFT, border=10)
        self.breakfest_count_combobox.Bind(wx.EVT_COMBOBOX, self.breakfest_total, self.breakfest_count_combobox)
        self.breakfest_count_combobox.Disable()
        self.breakfest_checkbox = wx.CheckBox(panel)
        main_sizer.Add(self.breakfest_checkbox, pos=(11, 2), flag=wx.ALL, border=5)
        self.breakfest_checkbox.Bind(wx.EVT_CHECKBOX, self.checkbox_breakfest, self.breakfest_checkbox)
        self.breakfest_checkbox.SetValue(False)

        breakfest_total_stat_txt = wx.StaticText(panel, label="Сніданки загально:")
        main_sizer.Add(breakfest_total_stat_txt, pos=(12, 0), flag=wx.LEFT, border=10)
        self.breakfest = wx.TextCtrl(panel)
        main_sizer.Add(self.breakfest, pos=(12, 1), flag=wx.EXPAND | wx.LEFT, border=10)
        self.breakfest.Disable()

        payment_type_stat_txt = wx.StaticText(panel, label="Тип розрахунку:")  # combobox
        main_sizer.Add(payment_type_stat_txt, pos=(13, 0), flag=wx.LEFT, border=10)
        payment_types = ["Готівка", "Кредитна карта", "Безготівковий переказ"]
        self.payment_type = wx.ComboBox(panel, choices=payment_types, style=wx.CB_READONLY)
        main_sizer.Add(self.payment_type, pos=(13, 1), flag=wx.EXPAND | wx.LEFT, border=10)
        self.payment_type.Bind(wx.EVT_COMBOBOX, self.payment_type_combobox)

        admin_name_stat_txt = wx.StaticText(panel, label="Ім'я адміністратора:")
        main_sizer.Add(admin_name_stat_txt, pos=(14, 0), flag=wx.LEFT, border=10)
        admins_names = ["Аліна", "Влад", "Сергій"]
        self.admin_name = wx.ComboBox(panel, choices=admins_names, style=wx.CB_READONLY)
        main_sizer.Add(self.admin_name, pos=(14, 1), flag=wx.EXPAND | wx.LEFT, border=10)
        self.admin_name.Bind(wx.EVT_COMBOBOX, self.administrator_name_combobox)
        self.admin_name.SetValue("Сергій")

        admin_surname_stat_txt = wx.StaticText(panel, label="Призвіще ініціали адміністратора:")
        main_sizer.Add(admin_surname_stat_txt, pos=(15, 0), flag=wx.LEFT, border=10)
        admins_surnames = ["Чубенко С.С.", "Тиченко В.С.", "Шевченко А.Є."]
        self.admin_surname_combobox = wx.ComboBox(panel, choices=admins_surnames, style=wx.CB_READONLY)
        main_sizer.Add(self.admin_surname_combobox, pos=(15, 1), flag=wx.EXPAND | wx.LEFT, border=10)
        self.admin_surname_combobox.Bind(wx.EVT_COMBOBOX, self.administrator_surname_combobox)
        self.admin_surname_combobox.SetValue("Чубенко С.С.")

        # without cash and acts

        numberofbill_stat_txt = wx.StaticText(panel, label="↓ Номер рахунку ↓")
        main_sizer.Add(numberofbill_stat_txt, pos=(1, 2), flag=wx.EXPAND | wx.LEFT | wx.CENTRE, border=10)
        self.numberofbill = wx.TextCtrl(panel)
        main_sizer.Add(self.numberofbill, pos=(2, 2), flag=wx.EXPAND | wx.LEFT, border=10)

        company_stat_txt = wx.StaticText(panel, label="↓ Назва компанії ↓")
        main_sizer.Add(company_stat_txt, pos=(3, 2), flag=wx.EXPAND | wx.LEFT, border=10)
        self.company = wx.TextCtrl(panel)
        main_sizer.Add(self.company, pos=(4, 2), flag=wx.EXPAND | wx.LEFT, border=10)

        compowner_stat_txt = wx.StaticText(panel, label="↓ Власник/Диерктор Компанії ↓")
        main_sizer.Add(compowner_stat_txt, pos=(5, 2), flag=wx.EXPAND | wx.LEFT, border=10)
        self.compowner = wx.TextCtrl(panel)
        main_sizer.Add(self.compowner, pos=(6, 2), flag=wx.EXPAND | wx.LEFT, border=10)

        comprequs_stat_txt = wx.StaticText(panel, label="↓ Реквізити компанії ↓")
        main_sizer.Add(comprequs_stat_txt, pos=(13, 2), flag=wx.EXPAND | wx.LEFT | wx.BOTTOM, border=10)
        self.comprequs = wx.TextCtrl(panel, style=wx.TE_MULTILINE)
        main_sizer.Add(self.comprequs, pos=(14, 2), span=(6, 2), flag=wx.EXPAND | wx.ALL, border=10)

        panel.SetSizer(main_sizer)

    # functions

    def make_date_changed(self, event):
        date_make = self.date_make.GetValue()
        date_make = date_make.Format("%d.%m.%Y")
        return date_make

    def checkin_date_changed(self, event):
        checkin_date = self.checkin_date.GetValue()
        checkin_date = checkin_date.Format("%d.%m.%Y")
        self.get_duration_accomodation()
        self.total_price_accomodation(wx.adv.EVT_DATE_CHANGED)
        if self.checkbox_tour_tax(wx.EVT_CHECKBOX) is True:
            self.tour_tax_total(wx.EVT_COMBOBOX)
        return checkin_date

    def checkout_date_changed(self, event):
        checkout_date = self.checkout_date.GetValue()
        checkout_date = checkout_date.Format("%d.%m.%Y")
        self.get_duration_accomodation()
        self.total_price_accomodation(wx.adv.EVT_DATE_CHANGED)
        if self.checkbox_tour_tax(wx.EVT_CHECKBOX) is True:
            self.tour_tax_total(wx.EVT_COMBOBOX)
        return checkout_date

    def get_duration_accomodation(self):
        date1 = self.checkin_date.GetValue()
        date2 = self.checkout_date.GetValue()
        delta = date2.Subtract(date1)
        duration_accomodation = delta.GetDays()
        return duration_accomodation

    def category_combobox(self, event):
        selected_category = self.category.GetValue()
        if selected_category:
            self.price_accomodation_PN_text_ctrl.SetValue(str(prices_default.prices[selected_category]))
            self.total_price_accomodation(wx.EVT_COMBOBOX)
        if self.checkbox_tour_tax(wx.EVT_CHECKBOX) is True:
            self.tour_tax_total(wx.EVT_COMBOBOX)
        return selected_category

    def price_pernight_changed(self, event):
        price_accomodation_PN = self.price_accomodation_PN_text_ctrl.GetValue()
        return price_accomodation_PN

    def total_price_accomodation(self, event):
        duration_accomodation = self.get_duration_accomodation()
        price_accomodation_PN = self.price_pernight_changed(wx.EVT_TEXT)
        count_of_rooms = self.count_of_rooms_combobox(wx.EVT_COMBOBOX)
        total_price_accomodation = Decimal(int(duration_accomodation) * float(price_accomodation_PN) *
                                           int(count_of_rooms)).quantize(TWOPLACES)
        self.total_price_accomodation_text_ctrl.SetValue(str(total_price_accomodation))
        return total_price_accomodation

    def checkbox_count_of_rooms(self, event):
        count_of_rooms_checbox = self.count_of_rooms_checkbox.GetValue()
        if count_of_rooms_checbox:
            self.count_of_rooms.Enable()
            self.count_of_rooms_combobox(event)
        else:
            self.count_of_rooms.SetValue(self.count_of_rooms_list[0])
            self.count_of_rooms_combobox(event)
            self.count_of_rooms.Disable()
            self.total_price_accomodation(wx.EVT_CHECKBOX)

    def count_of_guest_combobox(self, event):
        count_of_guest = int(self.count_of_guest.GetValue())
        return count_of_guest

    def count_of_rooms_combobox(self, event):
        count_of_rooms = int(self.count_of_rooms.GetValue())
        #
        if count_of_rooms > self.count_of_guest_combobox(wx.EVT_COMBOBOX):
            self.count_of_guest.SetValue(self.count_of_rooms.GetValue())
            self.tour_tax_total(wx.ALL)
        #
        return count_of_rooms

    def checkbox_tour_tax(self, event):
        tour_tax_checkbox = self.tour_tax_checkbox.GetValue()
        if tour_tax_checkbox:
            self.tour_tax_text_ctrl.Enable()
            return True
        else:
            self.tour_tax_text_ctrl.Disable()
            self.tour_tax_text_ctrl.SetValue("")
            return False

    def tour_tax_count(self, event):
        tour_tax_count = int(self.count_of_guest.GetValue()) * int(self.get_duration_accomodation())
        if self.checkbox_tour_tax(wx.EVT_CHECKBOX):
            return tour_tax_count
        else:
            return ""

    def tour_tax_total(self, event):
        if self.checkbox_tour_tax(wx.EVT_CHECKBOX):
            tour_tax_total = Decimal(int(self.tour_tax_count(wx.EVT_COMBOBOX)) *
                                     float(prices_default.prices[RoomType.TourTax])).quantize(TWOPLACES)
            self.tour_tax_text_ctrl.SetValue(str(tour_tax_total))
            return tour_tax_total
        else:
            return ""

    def tour_tax_confirm(self):
        return fr"Тур. збір: {self.tour_tax_total(wx.ALL)} грн." if self.checkbox_tour_tax(wx.EVT_CHECKBOX) else ""

    def checkbox_breakfest(self, event):
        checkbox_breakfest = self.breakfest_checkbox.GetValue()
        if checkbox_breakfest:
            self.breakfest.Enable()
            self.breakfest_count_combobox.Enable()
            self.breakfest_count_combobox.SetValue("1")
            self.breakfest_total(event)
        else:
            self.breakfest.SetValue("")
            self.breakfest_count_combobox.SetValue("")
            self.breakfest.Disable()
            self.breakfest_count_combobox.Disable()

    def breakfest_count(self):
        if self.breakfest_count_combobox.GetValue():
            breakfest_count = int(self.breakfest_count_combobox.GetValue())
            return breakfest_count
        else:
            return ""

    def breakfest_total(self, event):
        if self.breakfest_count_combobox.GetValue():
            breakfest_total = Decimal(int(self.breakfest_count()) *
                                      float(prices_default.prices[RoomType.Breakfest])).quantize(TWOPLACES)
            self.breakfest.SetValue(str(breakfest_total))
            return int(breakfest_total)
        else:
            return ""

    def breakfest_confirm(self):
        if self.breakfest_total(wx.EVT_TEXT):
            return fr"Сніданки {prices_default.prices[RoomType.Breakfest]} x {self.breakfest_count()}: {self.breakfest_total(wx.EVT_COMBOBOX)} грн."
        else:
            return ""

    def total_price(self):
        if self.breakfest_total(wx.EVT_COMBOBOX) and self.tour_tax_total(wx.EVT_COMBOBOX):
            total_price = Decimal(self.total_price_accomodation(wx.EVT_TEXT) + self.tour_tax_total(wx.EVT_TEXT) +
                                  self.breakfest_total(wx.EVT_COMBOBOX))
            return total_price
        elif self.tour_tax_total(wx.ALL):
            total_price = Decimal(self.total_price_accomodation(wx.EVT_TEXT) + self.tour_tax_total(wx.EVT_TEXT))
            return total_price
        elif self.breakfest_total(wx.EVT_COMBOBOX):
            total_price = Decimal(self.total_price_accomodation(wx.EVT_TEXT) + self.breakfest_total(wx.EVT_COMBOBOX))
            return total_price
        else:
            return Decimal(self.total_price_accomodation(wx.EVT_TEXT))

    def payment_type_combobox(self, event):
        payment_type = self.payment_type.GetValue()
        return payment_type

    def administrator_name_combobox(self, event):
        selected_admin_name = self.admin_name.GetValue()
        return selected_admin_name

    def administrator_surname_combobox(self, event):
        selected_admin_surname = self.admin_surname_combobox.GetValue()
        return selected_admin_surname

    # bill fillers
    def t_t_pos(self): return "2" if self.checkbox_tour_tax(wx.EVT_COMBOBOX) else ""

    def t_t_text(self): return "Туристичний збір" if self.checkbox_tour_tax(wx.EVT_COMBOBOX) else ""

    def t_t_price(self): return fr"{prices_default.prices[RoomType.TourTax]}" if self.checkbox_tour_tax(wx.EVT_COMBOBOX) else ""

    def br_pos(self): return "3" if self.breakfest_count() else ""

    def br_text(self): return "Сніданок" if self.breakfest_count() else ""

    def br_price(self): return fr"{prices_default.prices[RoomType.Breakfest]}" if self.breakfest_count() else ""

    # bill_wc and act fucns

    def services_count(self):
        counter = 1
        if self.tour_tax_text_ctrl.GetValue():
            counter += 1
        if self.breakfest.GetValue():
            counter += 1
        return counter

    # def company_name(self):
    #     return fr"{self.company}".upper()

    # makers
    def getdata(self):
        order_information = OrderInformation(
            name=self.guest_name_text_ctrl.GetValue(),

            datemake=self.make_date_changed(wx.adv.EVT_DATE_CHANGED),
            checkin=self.checkin_date_changed(wx.adv.EVT_DATE_CHANGED),
            date_checkout=self.checkout_date_changed(wx.adv.EVT_DATE_CHANGED),
            duration=self.get_duration_accomodation(),

            current_category=self.category_combobox(wx.EVT_COMBOBOX),
            price_per_night=self.price_accomodation_PN_text_ctrl.GetValue(),
            price_accomodation=str(self.total_price_accomodation(wx.ALL)),

            count_of_guests=str(self.count_of_rooms_combobox(wx.EVT_COMBOBOX)),
            count_of_rooms=str(self.count_of_rooms_combobox(wx.EVT_COMBOBOX)),

            t_t_pos=str(self.t_t_pos()),
            t_t_text=str(self.t_t_text()),
            tour_tax_count=str(self.tour_tax_count(wx.ALL)),
            tour_tax_price=str(self.t_t_price()),
            tour_tax_total=str(self.tour_tax_total(wx.ALL)),
            tourtaxconfirm=str(self.tour_tax_confirm()),

            br_pos=str(self.br_pos()),
            br_text=str(self.br_text()),
            brkfprice=str(self.br_price()),
            brkftcount=str(self.breakfest_count()),
            breakfest_total=str(self.breakfest_total(wx.EVT_COMBOBOX)),
            brkfstconfirm=str(self.breakfest_confirm()),

            total_price=str(self.total_price()),
            totalpriceend=str(self.total_price()),
            conftotprc=str(self.total_price()),

            payment_type=self.payment_type_combobox(wx.EVT_COMBOBOX),
            administrator=self.administrator_name_combobox(wx.EVT_COMBOBOX),
            admininnitials=self.administrator_surname_combobox(wx.EVT_COMBOBOX),

            numberofbill=self.numberofbill.GetValue(),
            company=self.company.GetValue(),
            compowner=self.compowner.GetValue(),
            comprequs=self.comprequs.GetValue()
        )
        return asdict(order_information)

    def make_confirm(self, event):

        confirm_form = docx.Document(fr"{curr_path}\resourses\confirm_form.docx")

        self.getdata()
        self.content = self.getdata()
        for j in self.content:
            for table in confirm_form.tables:
                for col in table.columns:
                    for cell in col.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                if run.text.find(j) >= 0:
                                    run.text = run.text.replace(j, str(self.content[j]))
                                    style = confirm_form.styles['Normal']
                                    font = style.font
                                    font.name = "Times New Roman"
                                    font.size = docx.shared.Pt(12)
        for table in confirm_form.tables:
            for row in table.rows:
                    if all(cell.text.isspace() or cell.text == '' for cell in row.cells):
                        table._element.remove(row._element)

        def saver_confirm():
            if not os.path.exists(fr"{curr_path}\Підтверження бронювання"):
                os.makedirs(fr"{curr_path}\Підтверження бронювання")
            confirm_form.save(
                fr"{curr_path}\Підтверження бронювання\Підтверження бронювання {self.guest_name_text_ctrl.GetValue()} {self.make_date_changed(wx.adv.EVT_DATE_CHANGED)}.docx")

            convert(
                fr"{curr_path}\Підтверження бронювання\Підтверження бронювання {self.guest_name_text_ctrl.GetValue()} {self.make_date_changed(wx.adv.EVT_DATE_CHANGED)}.docx",
                fr"{curr_path}\Підтверження бронювання\Підтверження бронювання {self.guest_name_text_ctrl.GetValue()} {self.make_date_changed(wx.adv.EVT_DATE_CHANGED)}.pdf")

        saver_confirm()



    def make_bill(self, event):

        bill_form = docx.Document(fr"{curr_path}\resourses\bill_form.docx")

        self.getdata()
        self.content = self.getdata()
        for j in self.content:
            for table in bill_form.tables:
                for col in table.columns:
                    for cell in col.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                if run.text.find(j) >= 0:
                                    run.text = run.text.replace(j, str(self.content[j]))
                                    style = bill_form.styles['Normal']
                                    font = style.font
                                    font.name = "Times New Roman"
                                    font.size = docx.shared.Pt(12)
        for table in bill_form.tables:
            for row in table.rows:
                if all(cell.text.isspace() or cell.text == '' for cell in row.cells):
                    table._element.remove(row._element)

        def saver_bill():
            if not os.path.exists(fr"{curr_path}\Рахунки"):
                os.makedirs(fr"{curr_path}\Рахунки")
            bill_form.save(
                fr"{curr_path}\Рахунки\Рахунок {self.guest_name_text_ctrl.GetValue()} {self.make_date_changed(wx.adv.EVT_DATE_CHANGED)}.docx")

            convert(
                fr"{curr_path}\Рахунки\Рахунок {self.guest_name_text_ctrl.GetValue()} {self.make_date_changed(wx.adv.EVT_DATE_CHANGED)}.docx",
                fr"{curr_path}\Рахунки\Рахунок {self.guest_name_text_ctrl.GetValue()} {self.make_date_changed(wx.adv.EVT_DATE_CHANGED)}.pdf")

        saver_bill()


    def make_bill_wc(self, event):
        bill_form_wc = load_workbook(fr"{curr_path}\resourses\bill_wc_form.xlsx")
        bill_form_wc_editor = bill_form_wc.active

        total_price_coins = f"{self.total_price():.2f}"[:-3:-1][::-1]
        unstndrt_digit = list(range(5, 21))
        stndrt_digit = ["2", "3", "4"]
        numbywords = num2words(int(self.total_price()), lang='uk')
        numbywords = numbywords.capitalize()

        def coin_marker():
            if str(self.total_price()).endswith(str(i for i in unstndrt_digit)):
                return "копійок"
            elif str(self.total_price()).endswith("1"):
                return "копійка"
            elif str(self.total_price()).endswith(str(i for i in stndrt_digit)):
                return "копійки"
            else:
                return "копійок"

        # out of table
        bill_form_wc_editor["C12"] = fr"Рахунок на оплату № {self.numberofbill.GetValue()} від {self.make_date_changed(wx.adv.EVT_DATE_CHANGED)} р."
        bill_form_wc_editor["H17"] = self.company.GetValue()
        bill_form_wc_editor["C30"] = fr"Всього найменувань {self.services_count()} на суму {int(self.total_price())} грн. {total_price_coins} {coin_marker()}"
        bill_form_wc_editor["C31"] = fr"{numbywords} грн. {total_price_coins} {coin_marker()}"

        # tables
        bill_form_wc_editor["E24"] = fr'Розміщення {self.guest_name_text_ctrl.GetValue()} з {self.checkin_date_changed(wx.adv.EVT_DATE_CHANGED)} по {self.checkout_date_changed(wx.adv.EVT_DATE_CHANGED)} у категорії "{self.category.GetValue()}"'
        bill_form_wc_editor["AC24"] = int(self.get_duration_accomodation())
        bill_form_wc_editor["AH24"] = float(self.price_accomodation_PN_text_ctrl.GetValue())
        bill_form_wc_editor["AK27"] = float(self.total_price())

        # breakfest
        if self.breakfest_checkbox.GetValue():
            bill_form_wc_editor["C25"] = 2
            bill_form_wc_editor["E25"] = fr"Сніданок для {self.guest_name_text_ctrl.GetValue()}"
            bill_form_wc_editor["AC25"] = int(self.breakfest_count())
            bill_form_wc_editor["AH25"] = float(prices_default.prices[RoomType.Breakfest])
            bill_form_wc_editor["AK25"] = float(self.breakfest_total(wx.EVT_COMBOBOX))

        # tour tax
        if self.tour_tax_checkbox.GetValue():
            bill_form_wc_editor["C26"] = 3
            bill_form_wc_editor["E26"] = "Туристичний збір"
            bill_form_wc_editor["AC26"] = int(self.tour_tax_count(wx.EVT_COMBOBOX))
            bill_form_wc_editor["AH26"] = float(prices_default.prices[RoomType.TourTax])
            bill_form_wc_editor["AK26"] = float(self.tour_tax_total(wx.EVT_COMBOBOX))

        if not self.breakfest_checkbox.GetValue():
            bill_form_wc_editor["C26"] = 2
            bill_form_wc_editor.delete_rows(25)
        if not self.tour_tax_checkbox.GetValue() and not self.breakfest_checkbox.GetValue():
            bill_form_wc_editor.delete_rows(25)
        elif not self.tour_tax_checkbox.GetValue():
            bill_form_wc_editor.delete_rows(26)

        if not os.path.exists(fr"{curr_path}\Безготівкові рахунки"):
            os.makedirs(fr"{curr_path}\Безготівкові рахунки")
        bill_form_wc.save(fr"{curr_path}\Безготівкові рахунки\Безготівковий рахунок №{self.numberofbill.GetValue()} {self.company.GetValue()} від {self.make_date_changed(wx.adv.EVT_DATE_CHANGED)}.xlsx")
        bill_form_wc.close()

    def make_act(self, event):
        act_form = load_workbook(fr"{curr_path}\resourses\act_form.xlsx")
        act_form_editor = act_form.active

        total_price_coins = f"{self.total_price():.2f}"[:-3:-1][::-1]
        unstndrt_digit = list(range(5, 21))
        stndrt_digit = ["2", "3", "4"]
        company_name = self.company.GetValue()

        def coin_marker():
            if str(self.total_price()).endswith(str(i for i in unstndrt_digit)):
                return "копійок"
            elif str(self.total_price()).endswith("1"):
                return "копійка"
            elif str(self.total_price()).endswith(str(i for i in stndrt_digit)):
                return "копійки"
            else:
                return "копійок"

        # out of table
        act_form_editor.cell(row=5, column=18, value=f"{self.company.GetValue()}")
        act_form_editor.cell(row=8, column=18, value=f"{self.compowner.GetValue()}")
        act_form_editor.cell(row=10, column=2, value=f"АКТ надання послуг\r№ {self.numberofbill.GetValue()} від {self.make_date_changed(wx.adv.EVT_DATE_CHANGED)} р.")
        act_form_editor.cell(row=12, column=2, value=f'Ми, що нижче підписалися, представник Замовника {company_name}, з одного боку, і представник\rВиконавця ТОВАРИСТВО З ОБМЕЖЕНОЮ ВІДПОВІДАЛЬНІСТЮ "ПАЛЕ РОЯЛЬ ОДЕСА" , з іншого\rбоку, склали цей акт про те, що на підставі наведених документів:')
        act_form_editor.cell(row=15, column=10, value=f"Рахунок на оплату № {self.numberofbill.GetValue()} від {self.make_date_changed(wx.adv.EVT_DATE_CHANGED)}")
        act_form_editor.cell(row=28, column=2, value=fr"Загальна вартість робіт (послуг) склала {int(self.total_price())} грн. {total_price_coins} {coin_marker()} без ПДВ")
        act_form_editor.cell(row=38, column=2, value=fr"{self.checkout_date_changed(wx.adv.EVT_DATE_CHANGED)}")
        act_form_editor.cell(row=38, column=18, value=fr"{self.checkout_date_changed(wx.adv.EVT_DATE_CHANGED)}")
        act_form_editor.cell(row=39, column=18, value=self.comprequs.GetValue())

        # tables
        act_form_editor.cell(row=21, column=4, value=fr'Розміщення {self.guest_name_text_ctrl.GetValue()} з {self.checkin_date_changed(wx.adv.EVT_DATE_CHANGED)} по {self.checkout_date_changed(wx.adv.EVT_DATE_CHANGED)} у категорії "{self.category.GetValue()}"')
        act_form_editor.cell(row=21, column=21, value=int(self.get_duration_accomodation()))
        act_form_editor.cell(row=21, column=26, value=float(self.price_accomodation_PN_text_ctrl.GetValue()))
        act_form_editor.cell(row=21, column=30, value=float(self.total_price()))

        # breakfest
        if self.breakfest_checkbox.GetValue():
            act_form_editor.cell(row=22, column=2, value=2)
            act_form_editor.cell(row=22, column=4, value=fr"Сніданок для {self.guest_name_text_ctrl.GetValue()}")
            act_form_editor.cell(row=22, column=21, value=int(self.breakfest_count()))
            act_form_editor.cell(row=22, column=26, value=float(prices_default.prices[RoomType.Breakfest]))
            act_form_editor.cell(row=22, column=30, value=float(self.breakfest_total(wx.EVT_COMBOBOX)))

        # tour tax
        if self.tour_tax_checkbox.GetValue():
            act_form_editor.cell(row=23, column=2, value=3)
            act_form_editor.cell(row=23, column=4, value="Туристичний збір")
            act_form_editor.cell(row=23, column=21, value=int(self.tour_tax_count(wx.EVT_COMBOBOX)))
            act_form_editor.cell(row=23, column=26, value=float(prices_default.prices[RoomType.TourTax]))
            act_form_editor.cell(row=23, column=30, value=float(self.tour_tax_total(wx.EVT_COMBOBOX)))

        if not self.breakfest_checkbox.GetValue():
            act_form_editor.cell(row=22, column=2, value=2)
            act_form_editor.cell(row=22, column=4, value="Туристичний збір")
            act_form_editor.cell(row=22, column=21, value=self.tour_tax_count(wx.EVT_COMBOBOX))
            act_form_editor.cell(row=22, column=26, value=prices_default.prices[RoomType.TourTax])
            act_form_editor.cell(row=22, column=30, value=self.tour_tax_total(wx.EVT_COMBOBOX))
            act_form_editor.cell(row=23, column=2, value="")
            act_form_editor.cell(row=23, column=4, value="")
            act_form_editor.cell(row=23, column=21, value="")
            act_form_editor.cell(row=23, column=26, value="")
            act_form_editor.cell(row=23, column=30, value="")
        if not self.tour_tax_checkbox.GetValue() and not self.breakfest_checkbox.GetValue():
            act_form_editor.cell(row=22, column=2, value="")
            act_form_editor.cell(row=22, column=4, value="")
            act_form_editor.cell(row=22, column=21, value="")
            act_form_editor.cell(row=22, column=26, value="")
            act_form_editor.cell(row=22, column=30, value="")
            act_form_editor.cell(row=23, column=2, value="")
            act_form_editor.cell(row=23, column=4, value="")
            act_form_editor.cell(row=23, column=21, value="")
            act_form_editor.cell(row=23, column=26, value="")
            act_form_editor.cell(row=23, column=30, value="")
        elif not self.tour_tax_checkbox.GetValue():
            act_form_editor.cell(row=23, column=2, value="")
            act_form_editor.cell(row=23, column=4, value="")
            act_form_editor.cell(row=23, column=21, value="")
            act_form_editor.cell(row=23, column=26, value="")
            act_form_editor.cell(row=23, column=30, value="")

        if not os.path.exists(fr"{curr_path}\Акти надання послуг"):
            os.makedirs(fr"{curr_path}\Акти надання послуг")
        act_form.save(fr"{curr_path}\Акти надання послуг\Акт надання послуг № {self.numberofbill.GetValue()} {company_name} від {self.make_date_changed(wx.adv.EVT_DATE_CHANGED)}.xlsx")
        act_form.close()

    # Setting price frame
    def show_settings_price_frame(self, event):
        setting_prie_frame = SettingPriceFrame(self, title="Налаштування цін за замовчанням")
        setting_prie_frame.SetSize(420, 310)
        setting_prie_frame.Centre()
        setting_prie_frame.Show()
        print("frame time")

    def onquit(self, event):
        self.Close()


# prices default editing and saving
class PricesDefault:

    def __init__(self, prices_dict):
        self.prices = prices_dict

    def update_config(self, key, value):
        self.prices[key] = value

    def save_to_file(self, filename):
        with open(filename, 'wb') as file:
            pickle.dump(self.prices, file)

    def load_from_file(self, filename):
        with open(filename, 'rb') as file:
            self.prices = pickle.load(file)


class SettingPriceFrame(wx.Frame):
    def __init__(self, parent, title):
        super().__init__(parent, title=title, style=wx.DEFAULT_FRAME_STYLE & ~wx.RESIZE_BORDER)
        prices_default.load_from_file("prices_default.pkl")

        spf_panel = wx.Panel(self)

        frame_sizer = wx.GridBagSizer(10, 10)

        top_text = wx.StaticText(spf_panel, label="Ціна за замовчанням")
        frame_sizer.Add(top_text, pos=(0, 1), flag=wx.EXPAND | wx.LEFT | wx.TOP | wx.RIGHT, border=10)

        standart_price_stat_txt = wx.StaticText(spf_panel, label="Стандарт ціна:")
        frame_sizer.Add(standart_price_stat_txt, pos=(1, 0), flag=wx.LEFT, border=10)
        self.standart_price_text_ctrl = wx.TextCtrl(spf_panel)
        self.standart_price_text_ctrl.SetValue(str(prices_default.prices[RoomType.Standart]))
        frame_sizer.Add(self.standart_price_text_ctrl, pos=(1, 1), flag=wx.EXPAND | wx.LEFT, border=10)

        classic_price_stat_txt = wx.StaticText(spf_panel, label="Класичний ціна:")
        frame_sizer.Add(classic_price_stat_txt, pos=(2, 0), flag=wx.LEFT, border=10)
        self.classic_price_text_ctrl = wx.TextCtrl(spf_panel)
        self.classic_price_text_ctrl.SetValue(str(prices_default.prices[RoomType.Classic]))
        frame_sizer.Add(self.classic_price_text_ctrl, pos=(2, 1), flag=wx.EXPAND | wx.LEFT, border=10)

        junior_suite_price_stat_txt = wx.StaticText(spf_panel, label="Напівлюкс ціна:")
        frame_sizer.Add(junior_suite_price_stat_txt, pos=(3, 0), flag=wx.LEFT, border=10)
        self.junior_suite_price_text_ctrl = wx.TextCtrl(spf_panel)
        self.junior_suite_price_text_ctrl.SetValue(str(prices_default.prices[RoomType.JuniorSuite]))
        frame_sizer.Add(self.junior_suite_price_text_ctrl, pos=(3, 1), flag=wx.EXPAND | wx.LEFT, border=10)

        suite_price_stat_txt = wx.StaticText(spf_panel, label="Люкс ціна:")
        frame_sizer.Add(suite_price_stat_txt, pos=(4, 0), flag=wx.LEFT, border=10)
        self.suite_price_text_ctrl = wx.TextCtrl(spf_panel)
        self.suite_price_text_ctrl.SetValue(str(prices_default.prices[RoomType.Suite]))
        frame_sizer.Add(self.suite_price_text_ctrl, pos=(4, 1), flag=wx.EXPAND | wx.LEFT, border=10)

        delux_price_stat_txt = wx.StaticText(spf_panel, label="ДеЛюкс ціна:")
        frame_sizer.Add(delux_price_stat_txt, pos=(5, 0), flag=wx.LEFT, border=10)
        self.delux_price_text_ctrl = wx.TextCtrl(spf_panel)
        self.delux_price_text_ctrl.SetValue(str(prices_default.prices[RoomType.DeLux]))
        frame_sizer.Add(self.delux_price_text_ctrl, pos=(5, 1), flag=wx.EXPAND | wx.LEFT, border=10)

        tourist_tax_price_stat_txt = wx.StaticText(spf_panel, label="Туристичний збір:")
        frame_sizer.Add(tourist_tax_price_stat_txt, pos=(6, 0), flag=wx.LEFT, border=10)
        self.tourist_tax_price_text_ctrl = wx.TextCtrl(spf_panel)
        self.tourist_tax_price_text_ctrl.SetValue(str(prices_default.prices[RoomType.TourTax]))
        frame_sizer.Add(self.tourist_tax_price_text_ctrl, pos=(6, 1), flag=wx.EXPAND | wx.LEFT, border=10)

        breakfest_price_stat_txt = wx.StaticText(spf_panel, label="Сніданок ціна:")
        frame_sizer.Add(breakfest_price_stat_txt, pos=(7, 0), flag=wx.LEFT, border=10)
        self.breakfest_price_text_ctrl = wx.TextCtrl(spf_panel)
        self.breakfest_price_text_ctrl.SetValue(str(prices_default.prices[RoomType.Breakfest]))
        frame_sizer.Add(self.breakfest_price_text_ctrl, pos=(7, 1), flag=wx.EXPAND | wx.LEFT, border=10)

        change_prices_button = wx.Button(spf_panel, label="Змінити\nціни\nза\nзамовчуванням")
        frame_sizer.Add(change_prices_button, pos=(0, 2), span=(9, 1), flag=wx.EXPAND | wx.ALL, border=10)
        change_prices_button.Bind(wx.EVT_BUTTON, self.change_prices, change_prices_button)

        spf_panel.SetSizer(frame_sizer)

        # functions

    def change_prices(self, event):

        prices_default.load_from_file("prices_default.pkl")

        for value in prices_default.prices:
            if value:
                prices_default.update_config(RoomType.Standart, self.standart_price_text_ctrl.GetValue())
                prices_default.update_config(RoomType.Classic, self.classic_price_text_ctrl.GetValue())
                prices_default.update_config(RoomType.JuniorSuite, self.junior_suite_price_text_ctrl.GetValue())
                prices_default.update_config(RoomType.Suite, self.suite_price_text_ctrl.GetValue())
                prices_default.update_config(RoomType.DeLux, self.delux_price_text_ctrl.GetValue())
                prices_default.update_config(RoomType.TourTax, self.tourist_tax_price_text_ctrl.GetValue())
                prices_default.update_config(RoomType.Breakfest, self.breakfest_price_text_ctrl.GetValue())
            else:
                value = Decimal("0.00")

        prices_default.save_to_file("prices_default.pkl")

        print("Change prices")


# information
class RoomType(str, Enum):
    Standart = "Стандарт"
    Classic = "Класичний"
    JuniorSuite = "Напівлюкс"
    Suite = "Люкс"
    DeLux = "ДеЛюкс"
    Breakfest = "Сніданок"
    TourTax = "Туристичний збір"


prices_default = PricesDefault({
    RoomType.Standart: Decimal("1000.00"),
    RoomType.Classic: Decimal("1200.00"),
    RoomType.JuniorSuite: Decimal("1500.00"),
    RoomType.Suite: Decimal("1800.00"),
    RoomType.DeLux: Decimal("2200.00"),
    RoomType.Breakfest: Decimal("190.00"),
    RoomType.TourTax: Decimal("33.50")
})


@dataclass
class OrderInformation:
    # confirm and bill data ↓
    name: str

    datemake: str
    checkin: str
    date_checkout: str
    duration: str

    current_category: str
    price_per_night: str
    price_accomodation: str

    count_of_guests: str
    count_of_rooms: str

    t_t_pos: str
    t_t_text: str
    tour_tax_price: str
    tour_tax_count: str
    tour_tax_total: str
    tourtaxconfirm: str

    br_pos: str
    br_text: str
    brkfprice: str
    brkftcount: str
    breakfest_total: str
    brkfstconfirm: str

    total_price: str
    totalpriceend: str
    conftotprc: str

    payment_type: str
    administrator: str
    admininnitials: str

    # bill_wc and acts data ↓
    numberofbill: str
    company: str
    compowner: str
    comprequs: str


# Main frame
main_frame = MyFrame(None, title="Admin_helper")
main_frame.SetSize(800, 700)
main_frame.Centre()
main_frame.Show()

app.MainLoop()
